import streamlit as st
import pandas as pd
import plotly.express as px
import requests
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from requests.exceptions import RequestException, ConnectionError, Timeout
from ai71 import AI71
import PyPDF2
import io
import random
import docx
import os
from docx import Document
from docx.shared import Inches
from datetime import datetime
import re
import logging
import base64
from typing import List, Dict, Any
import matplotlib.pyplot as plt
from bs4 import BeautifulSoup, NavigableString, Tag
from io import StringIO
import wikipedia
from typing import List, Optional
from httpx_sse import SSEError
from difflib import SequenceMatcher
from datetime import datetime
import spacy
import time
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import networkx as nx

nlp = spacy.load("en_core_web_sm")

# Error handling for optional dependencies
try:
    from streamlit_lottie import st_lottie
except ImportError:
    st.error("Missing dependency: streamlit_lottie. Please install it using 'pip install streamlit-lottie'")
    st.stop()

AI71_API_KEY = "AI71 Falcon API key"

# Initialize AI71 client
try:
    ai71 = AI71(AI71_API_KEY)
except Exception as e:
    st.error(f"Failed to initialize AI71 client: {str(e)}")
    st.stop()

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "uploaded_documents" not in st.session_state:
    st.session_state.uploaded_documents = []
if "case_precedents" not in st.session_state:
    st.session_state.case_precedents = []

def analyze_uploaded_document(file):
    content = ""
    if file.type == "application/pdf":
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            content += page.extract_text()
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file)
        for para in doc.paragraphs:
            content += para.text + "\n"
    else:
        content = file.getvalue().decode("utf-8")
    return content

def get_document_based_response(prompt, document_content):
    messages = [
        {"role": "system", "content": "You are a helpful legal assistant LexAI which has all the legal information in the world and is the the best assitand for lawyers, lawfirms and a common citizen. Answer questions based on the provided document content."},
        {"role": "user", "content": f"Document content: {document_content}\n\nQuestion: {prompt}"}
    ]
    try:
        completion = ai71.chat.completions.create(
            model="tiiuae/falcon-180b-chat",
            messages=messages,
            stream=False,
        )
        return completion.choices[0].message.content
    except Exception as e:
        return f"An error occurred while processing your request: {str(e)}"

def get_ai_response(prompt: str) -> str:
    """Gets the AI response based on the given prompt."""
    messages = [
        {"role": "system", "content": "You are a helpful legal assistant LexAI which has all the legal information in the world and is the the best assitand for lawyers, lawfirms and a common citizen, answer the question based on the US law but if the question lies out of the context of us law then answer it too saying i am LexAI and advanced legal assistant but this is what i know about the topic you are asking"},
        {"role": "user", "content": prompt}
    ]
    try:
        # First, try streaming
        response = ""
        for chunk in ai71.chat.completions.create(
            model="tiiuae/falcon-180b-chat",
            messages=messages,
            stream=True,
        ):
            if chunk.choices[0].delta.content:
                response += chunk.choices[0].delta.content
        return response
    except Exception as e:
        print(f"Streaming failed, falling back to non-streaming request. Error: {e}")
        try:
            # makes it fall back to non-streaming request
            completion = ai71.chat.completions.create(
                model="tiiuae/falcon-180b-chat",
                messages=messages,
                stream=False,
            )
            return completion.choices[0].message.content
        except Exception as e:
            print(f"An error occurred while getting AI response: {e}")
            return f"I apologize, but I encountered an error while processing your request. Error: {str(e)}"

def display_chat_history():
    for message in st.session_state.chat_history:
        if isinstance(message, tuple):
            if len(message) == 2:
                user_msg, bot_msg = message
                st.info(f"**You:** {user_msg}")
                st.success(f"**Bot:** {bot_msg}")
            else:
                st.error(f"Unexpected message format: {message}")
        elif isinstance(message, dict):
            if message.get('type') == 'wikipedia':
                st.success(f"**Bot:** Wikipedia Summary:\n{message.get('summary', 'No summary available.')}\n" +
                           (f"[Read more on Wikipedia]({message.get('url')})" if message.get('url') else ""))
            elif message.get('type') == 'web_search':
                web_results_msg = "Web Search Results:\n"
                for result in message.get('results', []):
                    web_results_msg += f"[{result.get('title', 'No title')}]({result.get('link', '#')})\n{result.get('snippet', 'No snippet available.')}\n\n"
                st.success(f"**Bot:** {web_results_msg}")
            else:
                st.error(f"Unknown message type: {message}")
        else:
            st.error(f"Unexpected message format: {message}")

def analyze_document(file) -> str:
    """Analyzes uploaded legal documents."""
    content = ""
    if file.type == "application/pdf":
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            content += page.extract_text()
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file)
        for para in doc.paragraphs:
            content += para.text + "\n"
    else:
        content = file.getvalue().decode("utf-8")
    
    return content[:5000]  # Limit content to 5000 characters for analysis

def search_web(query: str, num_results: int = 3) -> List[Dict[str, str]]:
    try:
        service = build("customsearch", "v1", developerKey="AIzaSyD-1OMuZ0CxGAek0PaXrzHOmcDWFvZQtm8")
        
        # Add legal-specific terms to the query
        legal_query = f"legal {query} law case precedent"
        
        # Execute the search request
        res = service.cse().list(q=legal_query, cx="877170db56f5c4629", num=num_results * 2).execute()
        
        results = []
        if "items" in res:
            for item in res["items"]:
                # Check if the result is relevant
                if any(keyword in item["title"].lower() or keyword in item["snippet"].lower() 
                       for keyword in ["law", "legal", "court", "case", "attorney", "lawyer"]):
                    result = {
                        "title": item["title"],
                        "link": item["link"],
                        "snippet": item["snippet"]
                    }
                    results.append(result)
                    if len(results) == num_results:
                        break
        
        return results
    except Exception as e:
        print(f"Error performing web search: {e}")
        return []

def perform_web_search(query: str) -> List[Dict[str, Any]]:
    """
    Performs a web search to find recent legal cost estimates.
    """
    url = f"https://www.google.com/search?q={query}"
    headers = {'User-Agent': 'Mozilla/5.0'}
    response = requests.get(url, headers=headers)
    soup = BeautifulSoup(response.content, 'html.parser')

    results = []
    for g in soup.find_all('div', class_='g'):
        anchors = g.find_all('a')
        if anchors:
            link = anchors[0]['href']
            title = g.find('h3', class_='r')
            if title:
                title = title.text
            else:
                title = "No title"
            snippet = g.find('div', class_='s')
            if snippet:
                snippet = snippet.text
            else:
                snippet = "No snippet"
            
            # Extract cost estimates from the snippet
            cost_estimates = extract_cost_estimates(snippet)
            
            if cost_estimates:
                results.append({
                    "title": title,
                    "link": link,
                    "cost_estimates": cost_estimates
                })

    return results[:3]  # Return top 3 results with their cost estimates

def comprehensive_document_analysis(content: str) -> Dict[str, Any]:
    """Performs a comprehensive analysis of the document, including web and Wikipedia searches."""
    try:
        analysis_prompt = f"Analyze the following legal document and provide a summary, potential issues, and key clauses:\n\n{content}"
        document_analysis = get_ai_response(analysis_prompt)
        
        # Extract main topics or keywords from the document
        topic_extraction_prompt = f"Extract the main topic or keyword from the following document summary:\n\n{document_analysis}"
        topics = get_ai_response(topic_extraction_prompt)
        
        web_results = search_web(topics)
        wiki_results = search_wikipedia(topics)
        
        return {
            "document_analysis": document_analysis,
            "related_articles": web_results or [],  # Ensure that this this is always a list
            "wikipedia_summary": wiki_results
        }
    except Exception as e:
        print(f"Error in comprehensive document analysis: {e}")
        return {
            "document_analysis": "Error occurred during analysis.",
            "related_articles": [],
            "wikipedia_summary": {"summary": "Error occurred", "url": "", "title": ""}
        }

def search_wikipedia(query: str, sentences: int = 2) -> Dict[str, str]:
    try:
        # Ensures that the query is a string before slicing
        truncated_query = str(query)[:300]
        
        # Search Wikipedia
        search_results = wikipedia.search(truncated_query, results=5)
        
        if not search_results:
            return {"summary": "No Wikipedia article found.", "url": "", "title": ""}
        
        # Find the most relevant page title
        best_match = max(search_results, key=lambda x: SequenceMatcher(None, truncated_query.lower(), x.lower()).ratio())
        
        try:
            page = wikipedia.page(best_match, auto_suggest=False)
            summary = wikipedia.summary(page.title, sentences=sentences, auto_suggest=False)
            return {"summary": summary, "url": page.url, "title": page.title}
        except wikipedia.exceptions.DisambiguationError as e:
            try:
                page = wikipedia.page(e.options[0], auto_suggest=False)
                summary = wikipedia.summary(page.title, sentences=sentences, auto_suggest=False)
                return {"summary": summary, "url": page.url, "title": page.title}
            except:
                pass
        except wikipedia.exceptions.PageError:
            pass
        
        # If no summary found after trying the best match and disambiguation
        return {"summary": "No relevant Wikipedia article found.", "url": "", "title": ""}
    except Exception as e:
        print(f"Error searching Wikipedia: {e}")
        return {"summary": f"Error searching Wikipedia: {str(e)}", "url": "", "title": ""}

def extract_important_info(text: str) -> str:
    """Extracts and highlights important information from the given text."""
    prompt = f"Extract and highlight the most important legal information from the following text. Use markdown to emphasize key points:\n\n{text}"
    return get_ai_response(prompt)

user_agents = [
    'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
    'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/14.0 Safari/605.1.15',
    'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:89.0) Gecko/20100101 Firefox/89.0',
    'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.101 Safari/537.36'
]

# Rate limiting parameters
MIN_DELAY = 3  # Minimum delay between requests in seconds
MAX_DELAY = 10  # Maximum delay between requests in seconds
last_request_time = 0

def get_random_user_agent():
    return random.choice(user_agents)

def rate_limit():
    global last_request_time
    current_time = time.time()
    time_since_last_request = current_time - last_request_time
    if time_since_last_request < MIN_DELAY:
        sleep_time = random.uniform(MIN_DELAY, MAX_DELAY)
        time.sleep(sleep_time)
    last_request_time = time.time()

def fetch_detailed_content(url):
    rate_limit()
    
    chrome_options = webdriver.ChromeOptions()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument(f"user-agent={get_random_user_agent()}")

    try:
        # Use webdriver_manager to handle driver installation
        service = Service(ChromeDriverManager().install())
        with webdriver.Chrome(service=service, options=chrome_options) as driver:
            driver.get(url)
            
            # Wait for the main content to load
            WebDriverWait(driver, 20).until(
                EC.presence_of_element_located((By.TAG_NAME, "body"))
            )

            # Scroll to load any lazy-loaded content
            driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            time.sleep(2)  # Wait for any dynamic content to load

            # Get the page source after JavaScript execution
            page_source = driver.page_source

            # Use BeautifulSoup for parsing
            soup = BeautifulSoup(page_source, 'html.parser')

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Extract main content (customize based on the website structure)
            main_content = soup.find('main') or soup.find('article') or soup.find('div', class_=re.compile('content|main'))

            if not main_content:
                main_content = soup.body

            # Extract text content
            text_content = main_content.get_text(separator='\n', strip=True)

            # Clean and process the content
            cleaned_content = clean_content(text_content)

            return cleaned_content

    except Exception as e:
        print(f"Error fetching content: {e}")
        return f"Unable to fetch detailed content. Error: {str(e)}", {}

def clean_content(text):
    # Remove extra whitespace and newlines
    text = re.sub(r'\s+', ' ', text).strip()
    
    # Remove any remaining HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    
    # Remove special characters and digits (customize as needed)
    text = re.sub(r'[^a-zA-Z\s.,;:?!-]', '', text)
    
    # Split into sentences
    sentences = re.split(r'(?<=[.!?])\s+', text)
    
    # Remove short sentences (likely to be noise)
    sentences = [s for s in sentences if len(s.split()) > 3]
    
    # Join sentences back together
    cleaned_text = ' '.join(sentences)
    
    return cleaned_text

def extract_structured_data(soup):
    structured_data = {}

    # Extract title
    title = soup.find('title')
    if title:
        structured_data['title'] = title.get_text(strip=True)

    # Extract meta description
    meta_desc = soup.find('meta', attrs={'name': 'description'})
    if meta_desc:
        structured_data['description'] = meta_desc.get('content', '')

    # Extract headings
    headings = []
    for tag in ['h1', 'h2', 'h3']:
        for heading in soup.find_all(tag):
            headings.append({
                'level': tag,
                'text': heading.get_text(strip=True)
            })
    structured_data['headings'] = headings

    # Extract links
    links = []
    for link in soup.find_all('a', href=True):
        links.append({
            'text': link.get_text(strip=True),
            'href': link['href']
        })
    structured_data['links'] = links

    # Extract images
    images = []
    for img in soup.find_all('img', src=True):
        images.append({
            'src': img['src'],
            'alt': img.get('alt', '')
        })
    structured_data['images'] = images

    return structured_data

def query_public_case_law(query: str) -> List[Dict[str, Any]]:
    """Query publicly available case law databases (Justia and CourtListener) to find related cases."""
    cases = []
    
    # Justia Search using Google
    justia_url = f"https://www.google.com/search?q={query}+case+law+site:law.justia.com"
    justia_headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
    }
    
    try:
        justia_response = requests.get(justia_url, headers=justia_headers)
        justia_response.raise_for_status()
        justia_soup = BeautifulSoup(justia_response.text, 'html.parser')
        
        justia_results = justia_soup.find_all('div', class_='g')
        
        for result in justia_results[:5]:  # Limits it to top 5 results
            title_elem = result.find('h3')
            link_elem = result.find('a')
            snippet_elem = result.find('div', class_='VwiC3b')
            
            if title_elem and link_elem and snippet_elem:
                title = title_elem.text
                link = link_elem['href']
                snippet = snippet_elem.text
                
                # it extract case name and citation from the title
                case_info = title.split(' - ')
                if len(case_info) >= 2:
                    case_name = case_info[0]
                    citation = case_info[1]
                else:
                    case_name = title
                    citation = "Citation not found"
                
                cases.append({
                    "source": "Justia",
                    "case_name": case_name,
                    "citation": citation,
                    "summary": snippet,
                    "url": link
                })
    except requests.RequestException as e:
        print(f"Error querying Justia: {e}")

    # CourtListener Search
    courtlistener_url = f"https://www.courtlistener.com/api/rest/v3/search/?q={query}&type=o&format=json"
    courtlistener_data = {}
    for attempt in range(3):  # Retry up to 3 times
        try:
            courtlistener_response = requests.get(courtlistener_url)
            courtlistener_response.raise_for_status()
            courtlistener_data = courtlistener_response.json()
            break
        except (requests.RequestException, ValueError) as e:
            print(f"Attempt {attempt + 1} failed: {e}")
            if attempt == 2:
                print(f"Failed to retrieve or parse data from CourtListener: {e}")
            time.sleep(2)

    if 'results' in courtlistener_data:
        for result in courtlistener_data['results'][:3]:  # Limit to 3 results
            case_url = f"https://www.courtlistener.com{result['absolute_url']}"
            cases.append({
                "source": "CourtListener",
                "case_name": result['caseName'],
                "date_filed": result['dateFiled'],
                "docket_number": result.get('docketNumber', 'Not available'),
                "court": result['court'],
                "url": case_url
            })

    return cases

def comprehensive_document_analysis(content: str) -> Dict[str, Any]:
    """Performs a comprehensive analysis of the document, including web and Wikipedia searches."""
    try:
        analysis_prompt = f"Analyze the following legal document and provide a summary, potential issues, and key clauses:\n\n{content}"
        document_analysis = get_ai_response(analysis_prompt)
        
        topic_extraction_prompt = f"Extract the main topics or keywords from the following document summary:\n\n{document_analysis}"
        topics = get_ai_response(topic_extraction_prompt)
        
        web_results = search_web(topics)
        wiki_results = search_wikipedia(topics)
        
        return {
            "document_analysis": document_analysis,
            "related_articles": web_results or [],
            "wikipedia_summary": wiki_results
        }
    except Exception as e:
        print(f"Error in comprehensive document analysis: {e}")
        return {
            "document_analysis": "Error occurred during analysis.",
            "related_articles": [],
            "wikipedia_summary": {"summary": "Error occurred", "url": "", "title": ""}
        }

def format_public_cases(cases: List[Dict[str, Any]]) -> str:
    """Format public cases for the AI prompt."""
    formatted = ""
    for case in cases:
        formatted += f"Source: {case['source']}\n"
        formatted += f"Case Name: {case['case_name']}\n"
        if 'citation' in case:
            formatted += f"Citation: {case['citation']}\n"
        if 'summary' in case:
            formatted += f"Summary: {case['summary']}\n"
        if 'date_filed' in case:
            formatted += f"Date Filed: {case['date_filed']}\n"
        if 'docket_number' in case:
            formatted += f"Docket Number: {case['docket_number']}\n"
        if 'court' in case:
            formatted += f"Court: {case['court']}\n"
        formatted += "\n"
    return formatted

def format_web_results(results: List[Dict[str, str]]) -> str:
    """Format web search results for the AI prompt."""
    formatted = ""
    for result in results:
        formatted += f"Title: {result['title']}\n"
        formatted += f"Snippet: {result['snippet']}\n"
        formatted += f"URL: {result['link']}\n\n"
    return formatted


def find_case_precedents(case_details: str) -> Dict[str, Any]:
    """Finds relevant case precedents based on provided details."""
    try:
        # Query public case law databases
        public_cases = query_public_case_law(case_details)

        # Perform web search
        web_results = search_web(f"legal precedent {case_details}", num_results=3)

        # Perform Wikipedia search
        wiki_result = search_wikipedia(f"legal case {case_details}")

        # Compile all information
        compilation_prompt = f"""
        Analyze the following case details and identify key legal concepts and relevant precedents,
        Analyze and the following case law information, focusing solely on factual elements and legal principles. Do not include any speculative or fictional content:

        Case Details: {case_details}

        Public Case Law Results:
        {format_public_cases(public_cases)}

        Web Search Results:
        {format_web_results(web_results)}

        Wikipedia Information:
        {wiki_result['summary']}

        Provide a well-structured summary highlighting the most relevant precedents and legal principles
        Do not introduce any hypothetical scenarios.
        """

        summary = get_ai_response(compilation_prompt)

        return {
            "summary": summary,
            "public_cases": public_cases,
            "web_results": web_results,
            "wikipedia": wiki_result
        }
    except Exception as e:
        print(f"An error occurred in find_case_precedents: {e}")
        return {
            "summary": f"An error occurred while finding case precedents: {str(e)}",
            "public_cases": [],
            "web_results": [],
            "wikipedia": {
                'title': 'Error',
                'summary': 'Unable to retrieve Wikipedia information',
                'url': ''
            }
        }

def safe_find(element, selector, class_=None, attr=None):
    """Safely find and extract text or attribute from an element."""
    found = element.find(selector, class_=class_) if class_ else element.find(selector)
    if found:
        return found.get(attr) if attr else found.text.strip()
    return "Not available"

def search_web_duckduckgo(query: str, num_results: int = 3, max_retries: int = 3) -> List[Dict[str, str]]:
    """
    Performs a web search using the Google Custom Search API.
    Returns a list of dictionaries containing search result title, link, and snippet.
    """
    api_key = "AIzaSyD-1OMuZ0CxGAek0PaXrzHOmcDWFvZQtm8"
    cse_id = "877170db56f5c4629"

    user_agents = [
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
        'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/14.0 Safari/605.1.15',
        'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.101 Safari/537.36'
    ]

    for attempt in range(max_retries):
        try:
            headers = {'User-Agent': random.choice(user_agents)}
            
            service = build("customsearch", "v1", developerKey=api_key)

            res = service.cse().list(q=query, cx=cse_id, num=num_results).execute()

            results = []
            if "items" in res:
                for item in res["items"]:
                    result = {
                        "title": item["title"],
                        "link": item["link"],
                        "snippet": item.get("snippet", "")
                    }
                    results.append(result)
                    if len(results) == num_results:
                        break

            return results

        except HttpError as e:
            print(f"HTTP error occurred: {e}. Attempt {attempt + 1} of {max_retries}")
        except ConnectionError as e:
            print(f"Connection error occurred: {e}. Attempt {attempt + 1} of {max_retries}")
        except Timeout as e:
            print(f"Timeout error occurred: {e}. Attempt {attempt + 1} of {max_retries}")
        except RequestException as e:
            print(f"An error occurred during the request: {e}. Attempt {attempt + 1} of {max_retries}")
        except Exception as e:
            print(f"An unexpected error occurred: {e}. Attempt {attempt + 1} of {max_retries}")

        # Exponential backoff
        time.sleep(2 ** attempt)

    print("Max retries reached. No results found.")
    return []

def estimate_legal_costs(case_type: str, complexity: str, state: str) -> Dict[str, Any]:
    """
    Estimates legal costs based on case type, complexity, and location.
    Performs web searches for more accurate estimates, lawyer recommendations, and similar cases.
    """
    base_costs = {
        "Simple": (150, 300),
        "Moderate": (250, 500),
        "Complex": (400, 1000)
    }
    
    case_type_multipliers = {
        "Civil Litigation": 1.2,
        "Criminal Law": 1.5,
        "Family Law": 1.0,
        "Business Law": 1.3,
        "Intellectual Property": 1.4,
        "Employment Law": 1.1,
        "Immigration Law": 1.0,
        "Real Estate Law": 1.2,
        "Personal Injury": 1.3,
        "Tax Law": 1.4,
    }
    
    estimated_hours = {
        "Simple": (10, 30),
        "Moderate": (30, 100),
        "Complex": (100, 300)
    }
    
    min_rate, max_rate = base_costs[complexity]
    
    multiplier = case_type_multipliers.get(case_type, 1.0)
    min_rate *= multiplier
    max_rate *= multiplier
    
    min_hours, max_hours = estimated_hours[complexity]
    min_total = min_rate * min_hours
    max_total = max_rate * max_hours
    
    cost_breakdown = {
        "Hourly rate range": f"${min_rate:.2f} - ${max_rate:.2f}",
        "Estimated hours": f"{min_hours} - {max_hours}",
        "Total cost range": f"${min_total:.2f} - ${max_total:.2f}",
    }
    
    search_query = f"{case_type} legal costs {state}"
    web_search_results = search_web_duckduckgo(search_query, num_results=3)
    
    high_cost_areas = [
        "Expert witnesses (especially in complex cases)",
        "Extensive document review and e-discovery",
        "Multiple depositions",
        "Prolonged trial periods",
        "Appeals process"
    ]
    
    cost_saving_tips = [
        "Consider alternative dispute resolution methods like mediation or arbitration",
        "Be organized and provide all relevant documents upfront to reduce billable hours",
        "Communicate efficiently with your lawyer, bundling questions when possible",
        "Ask for detailed invoices and review them carefully",
        "Discuss fee arrangements, such as flat fees or contingency fees, where applicable"
    ]
    
    lawyer_tips = [
        "Research and compare multiple lawyers or law firms",
        "Ask for references and read client reviews",
        "Discuss fee structures and payment plans upfront",
        "Consider lawyers with specific expertise in your case type",
        "Ensure clear communication and understanding of your case"
    ]

    return {
        "cost_breakdown": cost_breakdown,
        "high_cost_areas": high_cost_areas,
        "cost_saving_tips": cost_saving_tips,
        "finding_best_lawyer_tips": lawyer_tips,
        "web_search_results": web_search_results
    }

def extract_cost_estimates(text: str) -> List[str]:
    """
    Extracts cost estimates from the given text.
    """
    patterns = [
        r'\$\d{1,3}(?:,\d{3})*(?:\.\d{2})?',  # Matches currency amounts like $1,000.00
        r'\d{1,3}(?:,\d{3})*(?:\.\d{2})?\s*(?:USD|GBP|CAD|EUR)',  # Matches amounts with currency codes
        r'(?:USD|GBP|CAD|EUR)\s*\d{1,3}(?:,\d{3})*(?:\.\d{2})?'  # Matches currency codes before amounts
    ]
    
    estimates = []
    for pattern in patterns:
        matches = re.findall(pattern, text)
        estimates.extend(matches)
    
    return estimates

def legal_cost_estimator_ui():
    st.title("Legal Cost Estimator")
    
    case_types = [
        "Personal Injury", "Medical Malpractice", "Criminal Law", "Family Law",
        "Divorce", "Bankruptcy", "Business Law", "Employment Law",
        "Estate Planning", "Immigration Law", "Intellectual Property",
        "Real Estate Law", "Tax Law"
    ]
    case_type = st.selectbox("Select case type", case_types)
    
    complexity = st.selectbox("Select case complexity", ["Simple", "Moderate", "Complex"])
    
    states = [
        "Alabama", "Alaska", "Arizona", "Arkansas", "California", "Colorado", "Connecticut",
        "Delaware", "Florida", "Georgia", "Hawaii", "Idaho", "Illinois", "Indiana", "Iowa",
        "Kansas", "Kentucky", "Louisiana", "Maine", "Maryland", "Massachusetts", "Michigan",
        "Minnesota", "Mississippi", "Missouri", "Montana", "Nebraska", "Nevada", "New Hampshire",
        "New Jersey", "New Mexico", "New York", "North Carolina", "North Dakota", "Ohio",
        "Oklahoma", "Oregon", "Pennsylvania", "Rhode Island", "South Carolina", "South Dakota",
        "Tennessee", "Texas", "Utah", "Vermont", "Virginia", "Washington", "West Virginia",
        "Wisconsin", "Wyoming"
    ]
    state = st.selectbox("Select state", states)
    
    if st.button("Estimate Costs"):
        with st.spinner("Estimating costs and retrieving data..."):
            cost_estimate = estimate_legal_costs(case_type, complexity, state)
        
        st.header("Estimated Legal Costs")
        for key, value in cost_estimate["cost_breakdown"].items():
            st.write(f"**{key}:** {value}")
        
        st.header("Potential High-Cost Areas")
        for area in cost_estimate["high_cost_areas"]:
            st.write(f"- {area}")
        
        st.header("Cost-Saving Tips")
        for tip in cost_estimate["cost_saving_tips"]:
            st.write(f"- {tip}")
        
        st.header("Tips for Finding the Best Lawyer")
        for tip in cost_estimate["finding_best_lawyer_tips"]:
            st.write(f"- {tip}")
        
        st.header("Web Search Results")
        if cost_estimate["web_search_results"]:
            for result in cost_estimate["web_search_results"]:
                st.subheader(f"[{result['title']}]({result['link']})")
                st.write(result["snippet"])
                st.write("---")
        else:
            st.write("No web search results found for the selected criteria.")

def split_text(text, max_chunk_size=4000):
    return [text[i:i+max_chunk_size] for i in range(0, len(text), max_chunk_size)]

def analyze_contract(contract_text: str) -> Dict[str, Any]:
    """Analyzes the contract text for clauses, benefits, and potential exploits."""
    chunks = split_text(contract_text)
    full_analysis = ""

    for i, chunk in enumerate(chunks):
        analysis_prompt = f"""
        Analyze the following part of the contract ({i+1}/{len(chunks)}), identifying clauses that are favorable and unfavorable to each party involved. 
        Highlight potential areas of concern or clauses that could be exploited. 
        Provide specific examples within this part of the contract to support your analysis.

        **Contract Text (Part {i+1}/{len(chunks)}):**
        {chunk}
        """

        try:
            chunk_analysis = get_ai_response(analysis_prompt)
            full_analysis += chunk_analysis + "\n\n"
        except Exception as e:
            return {"error": f"Error analyzing part {i+1} of the contract: {str(e)}"}

    return {"analysis": full_analysis}

def contract_analysis_ui():
    st.subheader("Contract Analyzer")

    uploaded_file = st.file_uploader(
        "Upload a contract document (PDF, DOCX, or TXT)",
        type=["pdf", "docx", "txt"],
    )

    if uploaded_file:
        contract_text = analyze_uploaded_document(uploaded_file)

        if st.button("Analyze Contract"):
            with st.spinner("Analyzing contract..."):
                analysis_results = analyze_contract(contract_text)

            st.write("### Contract Analysis")
            if "error" in analysis_results:
                st.error(analysis_results["error"])
            else:
                st.write(analysis_results.get("analysis", "No analysis available."))

CASE_TYPES = [
    "Civil Rights", "Contract", "Real Property", "Tort", "Labor", "Intellectual Property",
    "Bankruptcy", "Immigration", "Tax", "Criminal", "Social Security", "Environmental"
]

DATA_SOURCES = {
    "Civil Rights": "https://www.uscourts.gov/statistics-reports/caseload-statistics-data-tables",
    "Contract": "https://www.uscourts.gov/statistics-reports/caseload-statistics-data-tables",
    "Real Property": "https://www.uscourts.gov/statistics-reports/caseload-statistics-data-tables",
    "Tort": "https://www.uscourts.gov/statistics-reports/caseload-statistics-data-tables",
    "Labor": "https://www.uscourts.gov/statistics-reports/caseload-statistics-data-tables",
    "Intellectual Property": "https://www.uscourts.gov/statistics-reports/caseload-statistics-data-tables",
    "Bankruptcy": "https://www.uscourts.gov/statistics-reports/caseload-statistics-data-tables",
    "Immigration": "https://www.uscourts.gov/statistics-reports/caseload-statistics-data-tables",
    "Tax": "https://www.uscourts.gov/statistics-reports/caseload-statistics-data-tables",
    "Criminal": "https://www.uscourts.gov/statistics-reports/caseload-statistics-data-tables",
    "Social Security": "https://www.uscourts.gov/statistics-reports/caseload-statistics-data-tables",
    "Environmental": "https://www.uscourts.gov/statistics-reports/caseload-statistics-data-tables"
}

def fetch_case_data(case_type: str) -> pd.DataFrame:
    """Fetches actual historical data for the given case type."""
    # This data is based on U.S. District Courts—Civil Cases Commenced, by Nature of Suit
    data = {
        "Civil Rights": [56422, 57040, 54847, 53499, 54012, 52850, 51739, 41520, 35793, 38033, 47209, 44637],
        "Contract": [31077, 29443, 28221, 28073, 28394, 29312, 28065, 26917, 28211, 30939, 36053, 35218],
        "Real Property": [13716, 12760, 12482, 12340, 12410, 12537, 12211, 13173, 13088, 13068, 12527, 11991],
        "Tort": [86690, 80331, 79235, 77630, 75007, 74708, 73785, 75275, 74240, 75309, 98437, 86129],
        "Labor": [19229, 18586, 19690, 18550, 17190, 17356, 18511, 18284, 17583, 21208, 21118, 18743],
        "Intellectual Property": [11971, 11307, 11920, 13215, 12304, 11576, 11195, 10526, 10577, 11349, 10636, 11475],
        "Bankruptcy": [47806, 47951, 47134, 46194, 39091, 38784, 38125, 37751, 37153, 43498, 41876, 45119],
        "Immigration": [6454, 6880, 9185, 8567, 9181, 8252, 7125, 7960, 8848, 9311, 8847, 7880],
        "Tax": [1486, 1235, 1265, 1205, 1412, 1350, 1219, 1148, 1107, 1216, 1096, 1139],
        "Criminal": [78864, 80897, 81374, 80069, 77357, 79787, 81553, 78127, 68856, 64565, 57287, 59453],
        "Social Security": [18271, 19811, 19276, 17452, 18193, 17988, 18502, 18831, 19220, 21310, 20506, 19185],
        "Environmental": [772, 1047, 1012, 1070, 1135, 1148, 993, 909, 1046, 1084, 894, 733]
    }
    
    df = pd.DataFrame({
        'Year': range(2011, 2023),
        'Number of Cases': data[case_type]
    })
    
    return df

def visualize_case_trends(case_type: str):
    """Visualizes case trends based on case type using actual historical data."""
    df = fetch_case_data(case_type)

    # Create a Plotly figure
    fig = px.line(df, x='Year', y='Number of Cases', title=f"Trend of {case_type} Cases (2011-2022)")
    fig.update_layout(
        xaxis_title="Year",
        yaxis_title="Number of Cases",
        hovermode="x unified"
    )
    fig.update_traces(mode="lines+markers")

    return fig, df


def case_trend_visualizer_ui():
    st.subheader("Case Trend Visualizer")

    st.warning("Please note that the data presented here is for U.S. federal courts. Data may vary slightly depending on the sources and reporting methods used.")

    case_type = st.selectbox("Select case type to visualize", CASE_TYPES)

    if 'current_case_type' not in st.session_state:
        st.session_state.current_case_type = case_type
    
    if 'current_data' not in st.session_state:
        st.session_state.current_data = None

    if st.button("Visualize Trend") or st.session_state.current_case_type != case_type:
        st.session_state.current_case_type = case_type
        with st.spinner("Fetching and visualizing data..."):
            fig, df = visualize_case_trends(case_type)
            st.session_state.current_data = df

            # Display the Plotly chart
            st.plotly_chart(fig, use_container_width=True)

            # Display Statistics
            st.subheader("Case Statistics")
            total_cases = df['Number of Cases'].sum()
            avg_cases = df['Number of Cases'].mean()
            max_year = df.loc[df['Number of Cases'].idxmax(), 'Year']
            min_year = df.loc[df['Number of Cases'].idxmin(), 'Year']

            col1, col2, col3 = st.columns(3)
            col1.metric("Total Cases (2011-2022)", f"{total_cases:,}")
            col2.metric("Average Cases per Year", f"{avg_cases:,.0f}")
            col3.metric("Peak Year", f"{max_year}")

            # Trend Description
            st.write("Trend Description:", get_trend_description(df))

    if st.session_state.current_data is not None:
        df = st.session_state.current_data

        # Interactive Analysis Section
        st.subheader("Interactive Analysis")

        # Year-over-Year Change
        df['YoY Change'] = df['Number of Cases'].pct_change() * 100
        yoy_fig = px.bar(df, x='Year', y='YoY Change', title="Year-over-Year Change in Case Numbers")
        st.plotly_chart(yoy_fig, use_container_width=True)

        # Moving Average with slider
        max_window = min(6, len(df))  # Ensure max window doesn't exceed data points
        window = st.slider("Select moving average window:", 2, max_window, 2)
        df['Moving Average'] = df['Number of Cases'].rolling(window=window).mean()

        # Create a new figure for the moving average
        ma_fig = px.line(df, x='Year', y=['Number of Cases', 'Moving Average'], title=f"{window}-Year Moving Average")
        st.plotly_chart(ma_fig, use_container_width=True)

        # Raw Data
        st.subheader("Raw Data")
        st.dataframe(df)

        # Download Options
        csv = df.to_csv(index=False)
        st.download_button(
            label="Download data as CSV",
            data=csv,
            file_name=f"{case_type.lower().replace(' ', '_')}_trend_data.csv",
            mime="text/csv",
        )

        # Additional Information & Data Sources
        st.subheader("Additional Information")
        info = get_additional_info(case_type)
        st.markdown(info)

        st.subheader("Data Sources")
        st.markdown(f"- [U.S. Courts Statistics & Reports]({DATA_SOURCES[case_type]})")

        # --- Web Search Results ---
        st.subheader("Web Search Results")
        search_query = f"{case_type} case trends legal data"
        web_results = search_web_duckduckgo(search_query, num_results=3)
        if web_results:
            for result in web_results:
                st.write(f"[{result['title']}]({result['link']})")
                st.write(f"{result['snippet']}")
                st.write("---")
        else:
            st.write("No relevant web search results found.")

def get_potential_factors(case_type):
    """Provide potential factors affecting the trend based on case type."""
    factors = {
        "Civil Rights": "Changes in social awareness, legislative reforms, or high-profile incidents.",
        "Contract": "Economic conditions, business climate, or changes in contract law.",
        "Real Property": "Housing market trends, zoning laws, or property rights issues.",
        "Tort": "Changes in liability laws, public awareness of rights, or notable precedent-setting cases.",
        "Labor": "Economic conditions, changes in labor laws, or shifts in employment practices.",
        "Intellectual Property": "Technological advancements, patent law changes, or increased digital content creation.",
        "Bankruptcy": "Economic recession, changes in bankruptcy laws, or financial market conditions.",
        "Immigration": "Changes in immigration policies, global events, or economic factors.",
        "Tax": "Tax law changes, economic conditions, or IRS enforcement priorities.",
        "Criminal": "Law enforcement practices, changes in criminal laws, or societal factors."
    }
    return factors.get(case_type, "Various legal, economic, and societal factors.")

def get_additional_info(case_type: str) -> str:
    """Provides additional information about the case type."""
    info = {
        "Civil Rights": """
        Civil Rights cases encompass a wide range of issues, including discrimination, voting rights, and civil liberties. 
        Key points:
        1. These cases often involve allegations of discrimination based on race, gender, age, disability, or other protected characteristics.
        2. The Civil Rights Act of 1964 is a cornerstone piece of legislation in many of these cases.
        3. There was a significant drop in cases from 2017 to 2018, possibly due to policy changes.
        4. A sharp increase occurred in 2020, likely influenced by social movements and high-profile incidents.
        5. The overall trend shows fluctuations, reflecting changing societal and political landscapes.
        6. Many civil rights cases are class action lawsuits, representing groups of individuals.
        7. These cases can involve both government entities and private organizations as defendants.
        8. The outcomes of civil rights cases often have far-reaching implications for societal norms and practices.
        9. Recent years have seen an increase in cases related to LGBTQ+ rights and protections.
        10. Civil rights cases related to technology and privacy issues are becoming more prevalent.
        11. The rise of social media has led to new types of civil rights cases involving online discrimination and harassment.
        12. Voting rights cases tend to spike around election years, particularly in contentious political climates.
        """,
        "Contract": """
        Contract cases involve disputes over agreements between parties.
        Key points:
        1. There's a general stability in the number of cases from 2011 to 2019.
        2. A noticeable increase occurred in 2020 and 2021, possibly due to COVID-19 related contract disputes.
        3. The trend suggests economic conditions and major events significantly impact contract litigation.
        4. Common types of contract disputes include breach of contract, contract interpretation, and enforcement of terms.
        5. B2B (Business-to-Business) contracts often form a significant portion of these cases.
        6. Employment contracts and non-compete agreements are frequent subjects of litigation.
        7. The rise of e-commerce has led to an increase in cases related to online contracts and terms of service.
        8. International contract disputes often involve complex jurisdictional issues.
        9. Alternative dispute resolution methods like arbitration are increasingly being used in contract cases.
        10. The Uniform Commercial Code (UCC) plays a crucial role in many contract disputes involving the sale of goods.
        11. Force majeure clauses have gained prominence in contract litigation, especially since the COVID-19 pandemic.
        12. Smart contracts and blockchain technology are introducing new complexities in contract law.
        """,
        "Real Property": """
        Real Property cases deal with land and property rights.
        Key points:
        1. The number of cases has remained relatively stable over the years.
        2. A slight increase is observed in 2018-2019, possibly due to changes in housing markets or property laws.
        3. The consistency in case numbers suggests enduring importance of property rights in legal disputes.
        4. Common issues include boundary disputes, easements, and zoning conflicts.
        5. Landlord-tenant disputes form a significant portion of real property cases.
        6. Foreclosure cases tend to increase during economic downturns.
        7. Environmental regulations increasingly impact real property law, leading to new types of cases.
        8. Cases involving homeowners' associations (HOAs) have become more common in recent years.
        9. Property tax disputes are a recurring theme in real property litigation.
        10. Eminent domain cases, while less frequent, often attract significant public attention.
        11. The rise of short-term rentals (e.g., Airbnb) has introduced new legal challenges in property law.
        12. Cases involving mineral rights and natural resource extraction remain important in certain regions.
        """,
        "Tort": """
        Tort cases involve civil wrongs that cause harm or loss.
        Key points:
        1. There's a general decline in tort cases from 2011 to 2019.
        2. A significant spike occurred in 2020, potentially related to the COVID-19 pandemic.
        3. The overall trend may reflect changes in liability laws and public awareness of legal rights.
        4. Personal injury cases, including car accidents and slip-and-falls, make up a large portion of tort litigation.
        5. Medical malpractice is a significant and often complex area of tort law.
        6. Product liability cases can lead to large class-action lawsuits against manufacturers.
        7. Defamation cases, including libel and slander, have evolved with the rise of social media.
        8. Environmental torts, such as cases related to pollution or toxic exposure, are increasingly common.
        9. Many states have implemented tort reform measures, affecting the number and nature of cases filed.
        10. Mass tort litigation, often involving pharmaceuticals or consumer products, can involve thousands of plaintiffs.
        11. Cybersecurity breaches have led to a new category of tort cases related to data privacy.
        12. The concept of 'loss of chance' in medical malpractice cases has gained traction in some jurisdictions.
        """,
        "Labor": """
        Labor cases involve disputes between employers and employees.
        Key points:
        1. The number of cases fluctuates year to year, reflecting changing labor market conditions.
        2. A notable increase occurred in 2019-2020, possibly due to pandemic-related employment issues.
        3. The trend highlights the ongoing importance of labor rights and workplace disputes.
        4. Wage and hour disputes, including overtime pay issues, are common in labor litigation.
        5. Discrimination and harassment cases form a significant portion of labor law disputes.
        6. Wrongful termination suits often spike during economic downturns.
        7. Cases involving employee classification (e.g., independent contractor vs. employee) have increased with the gig economy.
        8. Union-related disputes, while less common than in the past, still play a role in labor litigation.
        9. Workplace safety cases, including those related to OSHA regulations, are an important subset of labor law.
        10. The rise of remote work has introduced new legal questions in areas like workers' compensation.
        11. Non-compete and trade secret cases often intersect with labor law.
        12. Cases involving employee benefits and ERISA (Employee Retirement Income Security Act) are complex and frequent.
        """,
        "Intellectual Property": """
        Intellectual Property cases involve patents, copyrights, trademarks, and trade secrets.
        Key points:
        1. There's variability in the number of cases, with peaks in 2013 and 2019.
        2. The fluctuations may reflect changes in technology, innovation rates, and IP law developments.
        3. The overall trend underscores the critical role of IP in the modern, knowledge-based economy.
        4. Patent infringement cases, especially in the tech sector, often involve high stakes and complex technologies.
        5. Copyright cases have evolved with digital media, often involving issues of fair use and digital rights management.
        6. Trademark disputes frequently arise in e-commerce and social media contexts.
        7. Trade secret cases have gained prominence, particularly in industries with high employee mobility.
        8. The America Invents Act of 2011 significantly impacted patent litigation trends.
        9. International IP disputes often involve complex jurisdictional and enforcement issues.
        10. The rise of artificial intelligence has introduced new challenges in patent and copyright law.
        11. Design patent cases, especially in consumer products, have seen increased attention.
        12. IP cases in the pharmaceutical industry, including those related to generic drugs, remain highly impactful.
        """,
        "Bankruptcy": """
        Bankruptcy cases involve individuals or businesses seeking debt relief or reorganization.
        Key points:
        1. There's a general decline in bankruptcy cases from 2011 to 2019.
        2. A notable increase occurred in 2020, likely due to economic impacts of the COVID-19 pandemic.
        3. The trend reflects overall economic conditions and changes in bankruptcy laws.
        4. Chapter 7 (liquidation) and Chapter 13 (individual debt adjustment) are the most common types for individuals.
        5. Chapter 11 reorganizations, typically used by businesses, often attract significant media attention.
        6. The 2005 Bankruptcy Abuse Prevention and Consumer Protection Act significantly impacted filing trends.
        7. Student loan debt, while generally non-dischargeable, has become a major issue in bankruptcy discussions.
        8. Medical debt remains a leading cause of personal bankruptcy filings in the U.S.
        9. Cross-border insolvency cases have increased with globalization.
        10. The rise of cryptocurrency has introduced new complexities in bankruptcy proceedings.
        11. Small business bankruptcy rules were modified in 2020 to streamline the process.
        12. Bankruptcy filings often lag behind economic downturns, explaining delayed spikes in case numbers.
        """,
        "Immigration": """
        Immigration cases involve disputes over citizenship, deportation, and immigration status.
        Key points:
        1. There's significant variability in the number of cases, reflecting changing immigration policies.
        2. Peaks are observed in 2013 and 2019-2020, possibly due to policy changes and global events.
        3. The trend highlights the complex and evolving nature of immigration law and policy.
        4. Asylum cases form a significant portion of immigration litigation.
        5. Deportation and removal proceedings are among the most common types of immigration cases.
        6. Cases involving unaccompanied minors have gained prominence in recent years.
        7. Employment-based immigration disputes often involve visa status and labor certification issues.
        8. Family-based immigration cases, including marriage fraud investigations, remain common.
        9. The implementation and challenges to travel bans have led to spikes in certain types of cases.
        10. Naturalization application denials and delays have been subjects of increased litigation.
        11. Cases involving immigration detention conditions and practices have attracted public attention.
        12. The intersection of criminal law and immigration (crimmigration) has become an important area of focus.
        """,
        "Tax": """
        Tax cases involve disputes with tax authorities or challenges to tax laws.
        Key points:
        1. The number of tax cases has remained relatively stable over the years.
        2. Small fluctuations may reflect changes in tax laws or enforcement priorities.
        3. The consistent trend suggests ongoing importance of tax-related legal issues.
        4. Individual income tax disputes are the most common type of tax litigation.
        5. Corporate tax cases, while fewer in number, often involve higher monetary stakes.
        6. International tax issues, including transfer pricing disputes, have gained prominence.
        7. Tax fraud and evasion cases, though less frequent, attract significant attention and resources.
        8. Estate and gift tax disputes often involve complex valuations and family dynamics.
        9. Cases challenging the constitutionality of new tax laws or regulations occur periodically.
        10. Tax cases related to cryptocurrency and digital assets are an emerging area.
        11. Disputes over tax-exempt status for organizations have social and political implications.
        12. Cases involving tax credits and incentives, such as for renewable energy, form a specialized subset.
        """,
        "Criminal": """
        Criminal cases involve prosecutions for violations of criminal law.
        Key points:
        1. There's a general increase in criminal cases from 2011 to 2017.
        2. A significant decline is observed from 2018 to 2022.
        3. The trend may reflect changes in law enforcement priorities, criminal justice reform efforts, or reporting methods.
        4. Drug-related offenses consistently make up a large portion of federal criminal cases.
        5. White-collar crime prosecutions, including fraud and embezzlement, fluctuate with enforcement priorities.
        6. Immigration-related criminal cases have been significantly influenced by policy changes.
        7. Cybercrime prosecutions have increased with the rise of digital technologies.
        8. Terrorism-related cases, while relatively few in number, often involve complex investigations.
        9. Criminal justice reform efforts have impacted sentencing practices and case dispositions.
        10. The use of DNA evidence has influenced both new prosecutions and appeals of old convictions.
        11. Cases involving police conduct and qualified immunity have gained increased attention.
        12. The opioid crisis has led to a rise in both drug possession and distribution cases.
        """,
        "Social Security": """
        Social Security cases typically involve disputes over benefits or eligibility.
        Key points:
        1. The number of cases shows some variability, with a peak in 2019-2020.
        2. The trend may reflect changes in Social Security policies, demographic shifts, or economic conditions affecting benefit claims.
        3. Disability benefit denials and appeals form a large portion of Social Security cases.
        4. The aging of the baby boomer generation has influenced the volume and nature of cases.
        5. Cases often involve complex medical evidence and vocational assessments.
        6. The backlog of cases at the administrative level often impacts the number of court filings.
        7. Changes in the definition and evaluation of disabilities have affected case trends.
        8. Overpayment cases, where beneficiaries are asked to repay benefits, are a recurring issue.
        9. Cases involving the intersection of workers' compensation and Social Security benefits can be complex.
        10. The rise in mental health awareness has influenced disability claim patterns.
        11. Technological changes in case processing and evaluation have impacted trends.
        12. Cases involving Supplemental Security Income (SSI) often intersect with other public benefit programs.
        """,
        "Environmental": """
        Environmental cases involve disputes over environmental regulations, pollution, or natural resource management.
        Key points:
        1. The number of cases shows some variability, with peaks in 2015-2016.
        2. The trend may reflect changes in environmental policies, increased awareness of environmental issues, or specific environmental events or disasters.
        3. Clean Air Act and Clean Water Act violations are common subjects of litigation.
        4. Cases related to climate change have increased in recent years, often challenging government policies.
        5. Endangered Species Act cases often involve conflicts between conservation and development.
        6. Toxic tort cases, such as those involving lead contamination or industrial pollution, can be complex and long-lasting.
        7. Environmental impact assessment challenges are frequent in large development projects.
        8. Cases involving renewable energy projects and their environmental impacts have grown.
        9. Water rights disputes, particularly in drought-prone areas, form a significant subset of cases.
        10. Litigation over oil and gas drilling, including fracking, has been prominent in certain regions.
        11. Cases challenging or enforcing international environmental agreements are increasing.
        12. Environmental justice cases, addressing disproportionate environmental burdens on certain communities, have gained attention.
        """
    }
    return info.get(case_type, "Additional information not available for this case type.")

def get_trend_description(df):
    """Generate a description of the overall trend."""
    first_value = df['Number of Cases'].iloc[0]
    last_value = df['Number of Cases'].iloc[-1]
    if last_value > first_value:
        return "The number of cases has generally increased over the five-year period."
    elif last_value < first_value:
        return "The number of cases has generally decreased over the five-year period."
    else:
        return "The number of cases has remained relatively stable over the five-year period."


class LegalDataRetriever:
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
        })
        logging.basicConfig(level=logging.DEBUG)
        self.logger = logging.getLogger(__name__)

    def search_courtlistener(self, query: str) -> Dict[str, Any]:
        """
        Search CourtListener for case information.
        """
        url = f"https://www.courtlistener.com/api/rest/v3/search/?q={query}&type=o&format=json"
        for attempt in range(3):  # Retry up to 3 times
            try:
                response = self.session.get(url)
                response.raise_for_status()
                data = response.json()
                break
            except (requests.RequestException, ValueError) as e:
                self.logger.error(f"Attempt {attempt + 1} failed: {e}")
                if attempt == 2:
                    return {"error": f"Failed to retrieve or parse data from CourtListener: {e}"}
                time.sleep(2)  # Wait before retrying

        if data['count'] == 0:
            return {"error": "No results found"}

        result = data['results'][0]
        case_url = f"https://www.courtlistener.com{result['absolute_url']}"
        
        try:
            case_response = self.session.get(case_url)
            case_response.raise_for_status()
            soup = BeautifulSoup(case_response.text, 'html.parser')
        except requests.RequestException as e:
            self.logger.error(f"Failed to retrieve case page: {e}")
            return {"error": f"Failed to retrieve case page: {e}"}

        judges = self.extract_judges(soup)
        author = self.extract_author(soup, judges)
        court_opinion = self.extract_court_opinion(soup)

        return {
            "case_name": result['caseName'],
            "date_filed": result['dateFiled'],
            "docket_number": result.get('docketNumber', 'Not available'),
            "court": result['court'],
            "status": result.get('status', 'Not available'),
            "url": case_url,
            "judges": judges,
            "author": author,
            "court_opinion": court_opinion
        }

    def extract_judges(self, soup):
        judges = []
        judge_elements = soup.find_all('a', class_='judge-link')
        if judge_elements:
            judges = [judge.text.strip() for judge in judge_elements]
        else:
            judge_info = soup.find('p', class_='bottom')
            if judge_info:
                judges = [j.strip() for j in judge_info.text.split(',') if j.strip()]
        
        if not judges:
            self.logger.warning("No judges found in the HTML structure, searching in text content")
            text_content = soup.get_text()
            judge_patterns = [
                r'(?:Judge|Justice)[s]?:?\s*(.*?)\.',
                r'(?:Before|Authored by):?\s*(.*?)\.', 
                r'(.*?),\s*(?:Circuit Judge|District Judge|Chief Judge)'
            ]
            for pattern in judge_patterns:
                judge_match = re.search(pattern, text_content, re.IGNORECASE)
                if judge_match:
                    judges = [j.strip() for j in judge_match.group(1).split(',') if j.strip()]
                    break
        
        return judges if judges else ["Not available"]

    def extract_author(self, soup, judges):
        author = "Not available"
        author_elem = soup.find('span', class_='author')
        if author_elem:
            author = author_elem.text.strip()
        elif judges and judges[0] != "Not available":
            author = judges[0]
        
        if author == "Not available":
            self.logger.warning("No author found in the HTML structure, searching in text content")
            text_content = soup.get_text()
            author_patterns = [
                r'(?:Author|Written by):?\s*(.*?)\.', 
                r'(.*?)\s*delivered the opinion of the court',
                r'(.*?),\s*(?:Circuit Judge|District Judge|Chief Judge).*?writing for the court'
            ]
            for pattern in author_patterns:
                author_match = re.search(pattern, text_content, re.IGNORECASE)
                if author_match:
                    author = author_match.group(1).strip()
                    break
        
        return author

    def extract_court_opinion(self, soup):
        article_div = soup.find('article', class_='col-sm-9')
        if not article_div:
            self.logger.error("Could not find the main article div (col-sm-9).")
            return "Case details not available (main article div not found)."

        opinion_div = article_div.find('div', class_='tab-content')
        if not opinion_div:
            self.logger.error("Could not find the case details content (tab-content div).")
            return "Case details not available (tab-content div not found)."

        case_details = opinion_div.get_text(separator='\n', strip=True)

        # Clean up the text
        case_details = re.sub(r'\n+', '\n', case_details)
        case_details = re.sub(r'\s+', ' ', case_details)

        return case_details

    def search_justia(self, query: str) -> Dict[str, Any]:
        """
        Search Justia for case information.
        """
        url = f"https://law.justia.com/cases/?q={query}"
        response = self.session.get(url)
        
        if response.status_code != 200:
            return {"error": "Failed to retrieve data from Justia"}

        soup = BeautifulSoup(response.text, 'html.parser')
        results = soup.find_all('div', class_='case-listing')
        
        if not results:
            return {"error": "No results found"}

        first_result = results[0]
        return {
            "case_name": first_result.find('h6').text.strip(),
            "citation": first_result.find('p', class_='citation').text.strip(),
            "summary": first_result.find('p', class_='summary').text.strip(),
            "url": first_result.find('a')['href'],
        }

def case_info_retriever():
    st.subheader("Case Information Retriever")
    query = st.text_input("Enter case name, number, or any relevant information:")
    if st.button("Retrieve Case Information"):
        with st.spinner("Retrieving case information..."):
            result = get_case_information(query)
        
        if "error" in result:
            st.error(result["error"])
        else:
            st.success("Case information retrieved successfully!")
            
            # Display case information
            st.subheader("Case Details")
            col1, col2 = st.columns(2)
            with col1:
                st.write(f"**Case Name:** {result['case_name']}")
                st.write(f"**Date Filed:** {result['date_filed']}")
                st.write(f"**Docket Number:** {result['docket_number']}")
            with col2:
                st.write(f"**Court:** {result['court']}")
                st.write(f"**Status:** {result['status']}")
                st.write(f"**[View on CourtListener]({result['url']})**")
            
            # Display judges and author
            st.subheader("Judges and Author")
            st.write(f"**Judges:** {', '.join(result['judges'])}")
            st.write(f"**Author:** {result['author']}")
            
            # Display case details (formerly court opinion)
            st.subheader("Case Details")
            st.markdown(result['court_opinion'])
            
            # Option to download the case information
            case_info_text = f"""
            Case Name: {result['case_name']}
            Date Filed: {result['date_filed']}
            Docket Number: {result['docket_number']}
            Court: {result['court']}
            Status: {result['status']}
            Judges: {', '.join(result['judges'])}
            Author: {result['author']}
            
            Case Details:
            {result['court_opinion']}
            
            View on CourtListener: {result['url']}
            """
            
            st.download_button(
                label="Download Case Information",
                data=case_info_text,
                file_name="case_information.txt",
                mime="text/plain"
            )

def get_case_information(query: str) -> Dict[str, Any]:
    retriever = LegalDataRetriever()
    
    # Search CourtListener
    cl_info = retriever.search_courtlistener(query)
    if "error" not in cl_info:
        return cl_info
    
    # Search Justia if CourtListener fails
    justia_info = retriever.search_justia(query)
    if "error" not in justia_info:
        return justia_info
    
    return {"error": "Unable to find case information from available sources."}

def extract_text_from_document(uploaded_file):
    text = ""
    if uploaded_file.type == "application/pdf":
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        for page in pdf_reader.pages:
            text += page.extract_text()
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(uploaded_file)
        for para in doc.paragraphs:
            text += para.text + "\n"
    else:
        text = uploaded_file.getvalue().decode("utf-8")
    return text

def generate_legal_brief(case_info):
    chunks = split_text(case_info)
    full_brief = ""
    
    for i, chunk in enumerate(chunks):
        prompt = f"""Generate a part of a comprehensive legal brief based on the following information (generate point 8, 9 10 and 11 only if a case is provided where outcome is yet to come). This is part {i+1} of {len(chunks)}. Focus on:
        1. A summary of the facts
        2. Identification of key legal issues
        3. Relevant laws and precedents
        4. Legal analysis
        5. Conclusion and recommendations
        6. An analysis of why the winning party won
        7. A review of how the losing party could have potentially won
        8. How the user can win this case based on the provided information.
        9. Key areas where user should be carefull and could potentially loose this case.
        10. Relevant Arguments for this case to be provided in the court.
        11. predict wheter user can win this case or not.

        Case Information (Part {i+1}/{len(chunks)}):
        {chunk}

        Please provide a detailed and thorough response for the relevant sections based on this part of the information."""

        try:
            response = ai71.chat.completions.create(
                model="tiiuae/falcon-180b-chat",
                messages=[{"role": "user", "content": prompt}],
                stream=False,
            )
            full_brief += response.choices[0].message.content + "\n\n"
        except Exception as e:
            st.error(f"Error generating part {i+1} of the legal brief: {str(e)}")
            return "Unable to generate complete legal brief due to an error."
    
    return full_brief

def automated_legal_brief_generation_ui():
    st.title("Automated Legal Brief Generation")
    with st.expander("How to use"):
        st.write('''Enter the case details and based on that it will generate a legal brief and also provide you with the proper analysis of the case and how you can win this case and where you have to be carefull''')
    if 'legal_brief' not in st.session_state:
        st.session_state.legal_brief = ""

    input_method = st.radio("Choose input method:", ("Text Input", "Document Upload"))
    
    if input_method == "Text Input":
        case_info = st.text_area("Enter the case information:", height=300)
    else:
        uploaded_file = st.file_uploader("Upload a document containing case details (PDF, DOCX, or TXT)", type=["pdf", "docx", "txt"])
        if uploaded_file is not None:
            case_info = extract_text_from_document(uploaded_file)
        else:
            case_info = ""

    if st.button("Generate Legal Brief"):
        if case_info:
            with st.spinner("Generating comprehensive legal brief..."):
                st.session_state.legal_brief = generate_legal_brief(case_info)
            st.success("Legal brief generated successfully!")
        else:
            st.warning("Please provide case information to generate the brief.")

    if st.session_state.legal_brief:
        st.subheader("Generated Legal Brief")
        st.text_area("Legal Brief", st.session_state.legal_brief, height=400)
        
        st.download_button(
            label="Download Legal Brief",
            data=st.session_state.legal_brief,
            file_name="legal_brief.txt",
            mime="text/plain"
        )

PRACTICE_AREAS = [
    "Personal Injury", "Medical Malpractice", "Criminal Law", "DUI & DWI", "Family Law",
    "Divorce", "Bankruptcy", "Business Law", "Consumer Law", "Employment Law",
    "Estate Planning", "Foreclosure Defense", "Immigration Law", "Intellectual Property",
    "Nursing Home Abuse", "Probate", "Products Liability", "Real Estate Law", "Tax Law",
    "Traffic Tickets", "Workers' Compensation", "Agricultural Law", "Animal & Dog Law",
    "Antitrust Law", "Appeals & Appellate", "Arbitration & Mediation", "Asbestos & Mesothelioma",
    "Cannabis & Marijuana Law", "Civil Rights", "Collections", "Communications & Internet Law",
    "Construction Law", "Domestic Violence", "Education Law", "Elder Law",
    "Energy, Oil & Gas Law", "Entertainment & Sports Law", "Environmental Law",
    "Gov & Administrative Law", "Health Care Law", "Insurance Claims", "Insurance Defense",
    "International Law", "Juvenile Law", "Landlord Tenant", "Legal Malpractice",
    "Maritime Law", "Military Law", "Municipal Law", "Native American Law", "Patents",
    "Securities Law", "Social Security Disability", "Stockbroker & Investment Fraud",
    "Trademarks", "White Collar Crime"
]

STATES = [
    "Alabama", "Alaska", "Arizona", "Arkansas", "California", "Colorado", "Connecticut", "Delaware", "Florida", "Georgia",
    "Hawaii", "Idaho", "Illinois", "Indiana", "Iowa", "Kansas", "Kentucky", "Louisiana", "Maine", "Maryland",
    "Massachusetts", "Michigan", "Minnesota", "Mississippi", "Missouri", "Montana", "Nebraska", "Nevada", "New Hampshire", "New Jersey",
    "New Mexico", "New York", "North Carolina", "North Dakota", "Ohio", "Oklahoma", "Oregon", "Pennsylvania", "Rhode Island", "South Carolina",
    "South Dakota", "Tennessee", "Texas", "Utah", "Vermont", "Virginia", "Washington", "West Virginia", "Wisconsin", "Wyoming"
]

CITIES_BY_STATE = {
    "Alabama": ["Birmingham", "Montgomery", "Mobile", "Huntsville"],
    "Alaska": ["Anchorage", "Fairbanks", "Juneau"],
    "Arizona": ["Phoenix", "Tucson", "Mesa", "Chandler"],
    "Arkansas": ["Little Rock", "Fort Smith", "Fayetteville"],
    "California": ["Los Angeles", "San Francisco", "San Diego", "San Jose"],
    "Colorado": ["Denver", "Colorado Springs", "Aurora", "Fort Collins"],
    "Connecticut": ["Bridgeport", "New Haven", "Hartford", "Stamford"],
    "Delaware": ["Wilmington", "Dover", "Newark"],
    "Florida": ["Miami", "Orlando", "Jacksonville", "Tampa"],
    "Georgia": ["Atlanta", "Augusta", "Columbus", "Savannah"],
    "Hawaii": ["Honolulu", "Hilo", "Kailua"],
    "Idaho": ["Boise", "Nampa", "Meridian"],
    "Illinois": ["Chicago", "Aurora", "Rockford", "Joliet"],
    "Indiana": ["Indianapolis", "Fort Wayne", "Evansville"],
    "Iowa": ["Des Moines", "Cedar Rapids", "Davenport"],
    "Kansas": ["Wichita", "Overland Park", "Kansas City"],
    "Kentucky": ["Louisville", "Lexington", "Bowling Green"],
    "Louisiana": ["New Orleans", "Baton Rouge", "Shreveport"],
    "Maine": ["Portland", "Lewiston", "Bangor"],
    "Maryland": ["Baltimore", "Columbia", "Annapolis"],
    "Massachusetts": ["Boston", "Worcester", "Springfield"],
    "Michigan": ["Detroit", "Grand Rapids", "Ann Arbor"],
    "Minnesota": ["Minneapolis", "St. Paul", "Rochester"],
    "Mississippi": ["Jackson", "Gulfport", "Southaven"],
    "Missouri": ["Kansas City", "St. Louis", "Springfield"],
    "Montana": ["Billings", "Missoula", "Great Falls"],
    "Nebraska": ["Omaha", "Lincoln", "Bellevue"],
    "Nevada": ["Las Vegas", "Reno", "Henderson"],
    "New Hampshire": ["Manchester", "Nashua", "Concord"],
    "New Jersey": ["Newark", "Jersey City", "Paterson"],
    "New Mexico": ["Albuquerque", "Las Cruces", "Santa Fe"],
    "New York": ["New York City", "Buffalo", "Rochester", "Syracuse"],
    "North Carolina": ["Charlotte", "Raleigh", "Greensboro"],
    "North Dakota": ["Fargo", "Bismarck", "Grand Forks"],
    "Ohio": ["Columbus", "Cleveland", "Cincinnati"],
    "Oklahoma": ["Oklahoma City", "Tulsa", "Norman"],
    "Oregon": ["Portland", "Eugene", "Salem"],
    "Pennsylvania": ["Philadelphia", "Pittsburgh", "Allentown"],
    "Rhode Island": ["Providence", "Warwick", "Cranston"],
    "South Carolina": ["Charleston", "Columbia", "North Charleston"],
    "South Dakota": ["Sioux Falls", "Rapid City", "Aberdeen"],
    "Tennessee": ["Nashville", "Memphis", "Knoxville"],
    "Texas": ["Houston", "Dallas", "Austin", "San Antonio"],
    "Utah": ["Salt Lake City", "West Valley City", "Provo"],
    "Vermont": ["Burlington", "South Burlington", "Rutland"],
    "Virginia": ["Virginia Beach", "Norfolk", "Chesapeake"],
    "Washington": ["Seattle", "Spokane", "Tacoma"],
    "West Virginia": ["Charleston", "Huntington", "Morgantown"],
    "Wisconsin": ["Milwaukee", "Madison", "Green Bay"],
    "Wyoming": ["Cheyenne", "Casper", "Laramie"]
}

def find_lawyers(state, city=None, practice_area=None, pages=1):
    base_url = "https://www.justia.com/lawyers/"
    url = base_url
    
    if practice_area:
        url += f"{practice_area.lower().replace(' ', '-').replace('&', '-')}/"
    
    url += state.lower()
    
    if city:
        url += f"/{city.lower().replace(' ', '-')}"
    
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
    }
    
    names = []
    short_bios = []
    specializations = []
    universities = []
    addresses = []
    phones = []
    email_links = []
    site_links = []
    
    try:
        for page in range(1, pages + 1):
            page_url = f"{url}?page={page}"
            response = requests.get(page_url, headers=headers)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            results = soup.find_all('div', {'data-vars-action': 'OrganicListing'})
            
            for result in results:
                # Name
                try:
                    names.append(result.find('strong', {'class': 'lawyer-name'}).get_text().strip())
                except:
                    names.append('')
                
                # Short Bio
                try:
                    short_bios.append(result.find('div', {'class': 'lawyer-expl'}).get_text().strip())
                except:
                    short_bios.append('')
                
                # Specialization
                try:
                    specializations.append(result.find('span', {'class': '-practices'}).get_text().strip())
                except:
                    specializations.append('')
                
                # University
                try:
                    universities.append(result.find('span', {'class': '-law-schools'}).get_text().strip())
                except:
                    universities.append('')
                
                # Address
                try:
                    addresses.append(result.find('span', {'class': '-address'}).get_text().strip().replace("\t", '').replace('\n', ', '))
                except:
                    addresses.append('')
                
                # Phone
                try:
                    phones.append(result.find('strong', {'class': '-phone'}).get_text().strip())
                except:
                    phones.append('')
                
                # Email Link
                try:
                    email_links.append(result.find('a', {'class': '-email'}).get('href'))
                except:
                    email_links.append('')
                
                # Site Link
                try:
                    site_links.append(result.find('a', {'class': '-website'}).get('href'))
                except:
                    site_links.append('')
        
        df_lawyers = pd.DataFrame({
            'lawyer_name': names,
            'short_bio': short_bios,
            'specialization': specializations,
            'university': universities,
            'address': addresses,
            'phone': phones,
            'email_link': email_links,
            'site_link': site_links
        })
        
        return df_lawyers
    
    except requests.RequestException as e:
        st.error(f"An error occurred while fetching lawyer information: {str(e)}")
        return pd.DataFrame()

def lawyer_finder_ui():
    st.title("Find Lawyers in Your Area")
    
    col1, col2, col3 = st.columns(3)
    with col1:
        state = st.selectbox("Select a State:", STATES)
    
    with col2:
        cities = CITIES_BY_STATE.get(state, [])
        city = st.selectbox("Select a City:", cities)
    
    with col3:
        practice_area = st.selectbox("Select a Practice Area:", [""] + PRACTICE_AREAS)
    
    if not city:
        st.warning("Please select a city to continue.")
        return
    
    pages = st.slider("Number of pages", 1, 20, 1)
    
    if st.button("Find Lawyers", type="primary"):
        with st.spinner("Searching for lawyers in your area..."):
            df_lawyers = find_lawyers(state, city, practice_area, pages)
            
            if not df_lawyers.empty:
                st.success(f"Found {len(df_lawyers)} lawyers in {city}, {state}.")
                
                # Display results in a more visually appealing way
                st.subheader("Lawyer Directory")
                for i in range(0, len(df_lawyers), 3):
                    cols = st.columns(3)
                    for j in range(3):
                        if i + j < len(df_lawyers):
                            lawyer = df_lawyers.iloc[i + j]
                            with cols[j]:
                                st.markdown(f"**{lawyer['lawyer_name']}**")
                                st.markdown(f"*{lawyer['specialization']}*")
                                if lawyer['phone']:
                                    st.markdown(f"📞 {lawyer['phone']}")
                                if lawyer['email_link']:
                                    st.markdown(f"📧 [Email]({lawyer['email_link']})")
                                if lawyer['site_link']:
                                    st.markdown(f"🌐 [Website]({lawyer['site_link']})")
                                st.markdown("---")
                
                # Show CSV preview with vertical and horizontal scrolling
                st.subheader("Data Preview")
                st.dataframe(
                    df_lawyers,
                    height=400,
                    width=600,
                    use_container_width=True,
                )
                
                # Provide CSV download option
                csv = df_lawyers.to_csv(index=False)
                st.download_button(
                    label="Download complete data as CSV",
                    data=csv,
                    file_name="lawyers_data.csv",
                    mime="text/csv",
                )
            else:
                st.warning(f"No lawyers found in {city}, {state}. Try selecting a different city or state.")

# Streamlit App 
st.markdown("""
<style>
    .reportview-container {
        background: #f0f2f6;
    }
    .main .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
        padding-left: 5rem;
        padding-right: 5rem;
    }
    h1 {
        color: #3e6ef7;
    }
    h2 {
        color: #3B82F6;
    }
    .stButton>button {
        background-color: #3B82F6;
        color: white;
        border-radius: 5px;
    }
    .stTextInput>div>div>input {
        border-radius: 5px;
    }
</style>
""", unsafe_allow_html=True)

def load_lottieurl(url: str):
    try:
        r = requests.get(url)
        r.raise_for_status()  # Raises a HTTPError if the status is 4xx, 5xx
        return r.json()
    except requests.HTTPError as http_err:
        print(f"HTTP error occurred while loading Lottie animation: {http_err}")
    except requests.RequestException as req_err:
        print(f"Error occurred while loading Lottie animation: {req_err}")
    except ValueError as json_err:
        print(f"Error decoding JSON for Lottie animation: {json_err}")
    return None

# Streamlit App
st.title("Lex AI - Advanced Legal Assistant")

# Sidebar with feature selection
with st.sidebar:
    st.title("Lex AI")
    st.subheader("Advanced Legal Assistant")
    
    feature = st.selectbox(
        "Select a feature",
        ["Legal Chatbot", "Document Analysis", "Case Precedent Finder", "Legal Cost Estimator", "Contract Analysis", "Case Trend Visualizer", "Case Information Retrieval", "Automated Legal Brief Generation", "Find the Lawyers"]
    )
if feature == "Legal Chatbot":
    st.subheader("Legal Chatbot")
    
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    
    if 'uploaded_document' not in st.session_state:
        st.session_state.uploaded_document = None
    
    if 'chat_mode' not in st.session_state:
        st.session_state.chat_mode = "normal"
    
    # Document upload
    uploaded_file = st.file_uploader("Upload a legal document (PDF, DOCX, or TXT)", type=["pdf", "docx", "txt"])
    
    if uploaded_file:
        st.session_state.uploaded_document = analyze_uploaded_document(uploaded_file)
        st.success("Document uploaded successfully!")
    
    # Chat mode toggle
    if st.session_state.uploaded_document:
        if st.button("Switch Chat Mode"):
            st.session_state.chat_mode = "document" if st.session_state.chat_mode == "normal" else "normal"
        
        st.write(f"Current mode: {'Document-based' if st.session_state.chat_mode == 'document' else 'Normal'} chat")
    
    display_chat_history()
    
    user_input = st.text_input("Your legal question:")
    
    if user_input and st.button("Send"):
        with st.spinner("Processing your question..."):
            if st.session_state.chat_mode == "document" and st.session_state.uploaded_document:
                ai_response = get_document_based_response(user_input, st.session_state.uploaded_document)
                st.session_state.chat_history.append((user_input, ai_response))
            else:
                ai_response = get_ai_response(user_input)
                st.session_state.chat_history.append((user_input, ai_response))
                
                # Perform Wikipedia search
                wiki_result = search_wikipedia(user_input)
                st.session_state.chat_history.append({
                    'type': 'wikipedia',
                    'summary': wiki_result.get("summary", "No summary available."),
                    'url': wiki_result.get("url", "")
                })
                
                # Perform web search
                web_results = search_web_duckduckgo(user_input)
                st.session_state.chat_history.append({
                    'type': 'web_search',
                    'results': web_results
                })
        
        st.rerun()

elif feature == "Document Analysis":
    st.subheader("Legal Document Analyzer")
    
    uploaded_file = st.file_uploader("Upload a legal document (PDF, DOCX, or TXT)", type=["pdf", "docx", "txt"])
    
    if uploaded_file and st.button("Analyze Document"):
        with st.spinner("Analyzing document and gathering additional information..."):
            try:
                document_content = analyze_document(uploaded_file)
                analysis_results = comprehensive_document_analysis(document_content)
                
                st.write("Document Analysis:")
                st.write(analysis_results.get("document_analysis", "No analysis available."))
                
                st.write("Related Articles:")
                for article in analysis_results.get("related_articles", []):
                    st.write(f"- [{article.get('title', 'No title')}]({article.get('link', '#')})")
                    st.write(f"  {article.get('snippet', 'No snippet available.')}")
                
                st.write("Wikipedia Summary:")
                wiki_info = analysis_results.get("wikipedia_summary", {})
                st.write(f"**{wiki_info.get('title', 'No title')}**")
                st.write(wiki_info.get('summary', 'No summary available.'))
                if wiki_info.get('url'):
                    st.write(f"[Read more on Wikipedia]({wiki_info['url']})")
            except Exception as e:
                st.error(f"An error occurred during document analysis: {str(e)}")

elif feature == "Case Precedent Finder":
    st.subheader("Case Precedent Finder")
    if 'precedents' not in st.session_state:
        st.session_state.precedents = None
    
    case_details = st.text_area("Enter case details:", height=100)
    if st.button("Find Precedents", type="primary"):
        with st.spinner("Searching for relevant case precedents..."):
            try:
                st.session_state.precedents = find_case_precedents(case_details)
            except Exception as e:
                st.error(f"An error occurred while finding case precedents: {str(e)}")
    
    if st.session_state.precedents:
        precedents = st.session_state.precedents
        
        st.markdown("## Summary of Relevant Case Precedents")
        st.info(precedents["summary"])
        
        st.markdown("## Related Cases from Public Databases")
        for i, case in enumerate(precedents["public_cases"], 1):
            st.markdown(f"### {i}. {case['case_name']}")
            col1, col2 = st.columns([1, 2])
            with col1:
                st.markdown(f"**Source:** {case['source']}")
                st.markdown(f"**URL:** [View Case]({case['url']})")
            with col2:
                for field in ['summary', 'date_filed', 'docket_number', 'court']:
                    if field in case and case[field]:
                        st.markdown(f"**{field.replace('_', ' ').title()}:** {case[field]}")
            st.markdown("---")
        
        st.markdown("## Additional Web Results")
        for i, result in enumerate(precedents["web_results"], 1):
            st.markdown(f"### {i}. {result['title']}")
            st.markdown(f"**Source:** [{result['link']}]({result['link']})")
            st.markdown(f"**Snippet:** {result['snippet']}")
            st.markdown("---")
        
        if precedents["wikipedia"]:
            st.markdown("## Wikipedia Information")
            wiki_info = precedents["wikipedia"]
            st.markdown(f"### {wiki_info['title']}")
            st.markdown(wiki_info['summary'])
            st.markdown(f"[Read more on Wikipedia]({wiki_info['url']})")
        st.markdown(
        """
        <style>
            .stTextArea > div > div > textarea {
                font-size: 16px;
            }
            h1, h2, h3 {
                margin-top: 1em;
                margin-bottom: 0.5em;
            }
        </style>
        """,
        unsafe_allow_html=True
    )

elif feature == "Legal Cost Estimator":
    legal_cost_estimator_ui()

elif feature == "Contract Analysis":
    contract_analysis_ui()

elif feature == "Case Trend Visualizer":
    case_trend_visualizer_ui()

elif feature == "Case Information Retrieval":
    case_info_retriever()

elif feature == "Automated Legal Brief Generation":
    automated_legal_brief_generation_ui()

elif feature == "Find the Lawyers":
    lawyer_finder_ui()

st.markdown("---")
st.markdown(
    """
    <div style="text-align: center;">
        <p>© 2023 Lex AI. All rights reserved.</p>
        <p><small>Disclaimer: This tool provides general legal information and assistance. It is not a substitute for professional legal advice. Please consult with a qualified attorney for specific legal matters.</small></p>
    </div>
    """,
    unsafe_allow_html=True
)

if __name__ == "__main__":
    st.sidebar.info("Select a feature from the dropdown above to get started.")
